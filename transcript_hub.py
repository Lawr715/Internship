import os
import re
import json
import csv
from io import StringIO, BytesIO
from datetime import datetime
from collections import Counter

import streamlit as st
# import streamlit.runtime.scriptrunner.script_runner
import matplotlib.pyplot as plt
from docx import Document
from keybert import KeyBERT
from transformers import pipeline

# Streamlit file watcher workaround
# os.environ["STREAMLIT_WATCHER_TYPE"] = "none"
# streamlit.runtime.scriptrunner.script_runner.ScriptRunner._on_file_change = lambda self, changed: None

# Helper to read .docx content
def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def generate_docx(transcript):
    doc = Document()
    doc.add_heading(f"{transcript['type']} – {transcript['location']} – {transcript['date']}", level=1)
    doc.add_paragraph("Summary:")
    doc.add_paragraph(transcript.get("summary", ""))
    doc.add_paragraph("Top Themes:")
    doc.add_paragraph(", ".join(transcript.get("themes", [])))
    doc.add_paragraph("Transcript:")
    doc.add_paragraph(transcript["text"])

    if transcript["tags"]:
        doc.add_paragraph("Tags: " + ", ".join(transcript["tags"]))
    if transcript["comments"]:
        doc.add_paragraph("Comments:")
        for comment in transcript["comments"]:
            doc.add_paragraph(f"- {comment}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Initialize session state for transcripts
if "transcripts" not in st.session_state:
    st.session_state.transcripts = []
if "kw_model" not in st.session_state:
    st.session_state.kw_model = KeyBERT("all-MiniLM-L6-v2")
if "summarizer" not in st.session_state:
    st.session_state.summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

# Password protection
PASSWORD = "transcripthub2025"
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    password = st.text_input("Enter password", type="password")
    if password == PASSWORD:
        st.session_state.authenticated = True
        st.rerun()
    else:
        st.stop()

st.set_page_config(page_title="Transcript Hub", layout="wide")
st.title("🧰 Transcript Hub")
st.caption("Internal tool to manage and analyze finalized KII and FGD transcripts")

# Upload section
st.header("📤 Upload Transcript")

uploaded_file = st.file_uploader("Upload a .txt or .docx file", type=["txt", "docx"])
text = ""

if uploaded_file:
    if uploaded_file.type == "text/plain":
        text = StringIO(uploaded_file.getvalue().decode("utf-8")).read()
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = extract_text_from_docx(uploaded_file)

# Collect existing values for suggestions
existing_locations = sorted({t["location"] for t in st.session_state.transcripts if t["location"]})
existing_languages = sorted({t["language"] for t in st.session_state.transcripts if t["language"]})

col1, col2, col3, col4 = st.columns(4)
with col1:
    location_option = st.selectbox("Research Site", existing_locations + ["Other"])
    if location_option == "Other":
        location = st.text_input("Enter new Research Site")
    else:
        location = location_option

with col2:
    t_type = st.selectbox("Type", ["FGD", "KII"])

with col3:
    language_option = st.selectbox("Language", existing_languages + ["Other"])
    if language_option == "Other":
        language = st.text_input("Enter new Language")
    else:
        language = language_option

with col4:
    date_str = st.date_input("Date", datetime.now()).strftime("%Y-%m-%d")

if st.button("Upload Transcript") and text:
    summary = st.session_state.summarizer(text[:1024], max_length=100, min_length=30, do_sample=False)[0]["summary_text"]
    st.session_state.transcripts.insert(0, {
        "text": text,
        "location": location,
        "type": t_type,
        "language": language,
        "date": date_str,
        "tags": [],
        "comments": [],
        "themes": [],
        "summary": summary
    })
    st.success("Transcript uploaded successfully!")

# Auto-Extracted Themes
st.header("🧠 Auto-Extracted Themes")

delete_index = None  # <-- Add this flag before the loop

for idx, t in enumerate(st.session_state.transcripts):
    if "themes" not in t or not t["themes"]:
        keywords = st.session_state.kw_model.extract_keywords(t["text"], top_n=5, stop_words="english")
        t["themes"] = [kw[0] for kw in keywords]

    with st.expander(f"{t['date']} • {t['type']} • {t['location']} • {t['language'] or 'Unknown Language'}"):
        st.markdown(f"**Top Themes:** {', '.join(t['themes'])}")
        st.markdown("**Summary:**")
        st.info(t.get("summary", "No summary available."))

        # --- Per-transcript export buttons ---
        json_data = json.dumps(t, indent=2)
        st.download_button(
            "📄 Download Transcript JSON",
            data=json_data,
            file_name=f"{t['type']}_{t['location']}_{t['date']}.json",
            mime="application/json"
        )
        st.download_button(
            "📝 Download Plain Text",
            data=t['text'],
            file_name=f"{t['type']}_{t['location']}_{t['date']}.txt",
            mime="text/plain"
        )
        docx_file = generate_docx(t)
        st.download_button(
            "📄 Download DOCX",
            data=docx_file,
            file_name=f"{t['type']}_{t['location']}_{t['date']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        # --- End per-transcript export buttons ---

        if st.button(f"Re-run extraction (transcript {idx+1})"):
            keywords = st.session_state.kw_model.extract_keywords(t["text"], top_n=5, stop_words="english")
            t["themes"] = [kw[0] for kw in keywords]
            st.success("Themes re-extracted.")

        if st.button(f"Re-run summarization (transcript {idx+1})"):
            t["summary"] = st.session_state.summarizer(t["text"][:1024], max_length=100, min_length=30, do_sample=False)[0]["summary_text"]
            st.success("Summary updated.")

        st.markdown("### ✏️ Edit Transcript Info")
        t["location"] = st.text_input(f"Edit Location (transcript {idx+1})", t["location"], key=f"edit_loc_{idx}")
        t["type"] = st.selectbox(f"Edit Type (transcript {idx+1})", ["FGD", "KII"], index=["FGD", "KII"].index(t["type"]), key=f"edit_type_{idx}")
        t["language"] = st.text_input(f"Edit Language (transcript {idx+1})", t["language"], key=f"edit_lang_{idx}")
        t["date"] = st.date_input(f"Edit Date (transcript {idx+1})", datetime.strptime(t["date"], "%Y-%m-%d"), key=f"edit_date_{idx}").strftime("%Y-%m-%d")

        if st.button(f"🗑 Delete transcript {idx+1}"):
            delete_index = idx

# After the for loop, handle deletion and rerun
if delete_index is not None:
    st.session_state.transcripts.pop(delete_index)
    st.experimental_rerun()

# Visualize term frequency across all transcripts
theme_counter = Counter()
for t in st.session_state.transcripts:
    theme_counter.update(t["themes"])

if theme_counter:
    st.subheader("📊 Most Frequent Themes")
    top_themes = theme_counter.most_common(10)
    labels, values = zip(*top_themes)
    fig, ax = plt.subplots()
    ax.barh(labels, values)
    ax.invert_yaxis()
    st.pyplot(fig)

# Sidebar Filters
st.sidebar.header("📂 Filter Transcripts")
sites = sorted(list({t["location"] for t in st.session_state.transcripts}))
types = sorted(list({t["type"] for t in st.session_state.transcripts}))
languages = sorted(list({t["language"] for t in st.session_state.transcripts}))

selected_site = st.sidebar.selectbox("Site", ["All"] + sites)
selected_type = st.sidebar.selectbox("Type", ["All"] + types)
selected_lang = st.sidebar.selectbox("Language", ["All"] + languages)

# --- Session Save/Load ---
if st.sidebar.button("💾 Save Session"):
    session_json = json.dumps(st.session_state.transcripts, indent=2)
    st.sidebar.download_button("Download Session JSON", session_json, file_name="session_backup.json", mime="application/json")

uploaded_session = st.sidebar.file_uploader("📂 Load Session", type=["json"])
if uploaded_session:
    session_data = json.load(uploaded_session)
    st.session_state.transcripts = session_data
    st.success("Session restored!")
    st.experimental_rerun()
# --- End Session Save/Load ---

# Search section
st.header("🔎 Search Transcripts")
search_term = st.text_input("Search by keyword or phrase (e.g., discipline, barangay program)")

def extract_context(text, keyword, context_lines=2):
    lines = text.split("\n")
    keyword_lower = keyword.lower()
    matches = []
    for i, line in enumerate(lines):
        if keyword_lower in line.lower():
            start = max(i - context_lines, 0)
            end = min(i + context_lines + 1, len(lines))
            snippet = "\n".join(lines[start:end])
            matches.append(snippet)
    return matches

# Tag/Theme Summary Panel
st.header("🏷️ Tag and Theme Overview")

# Aggregate all tags and themes
all_tags = []
all_themes = []
for t in st.session_state.transcripts:
    all_tags.extend(t["tags"])
    all_themes.extend(t["themes"])

tag_counts = Counter(all_tags)
theme_counts = Counter(all_themes)

if tag_counts:
    st.subheader("Top Tags")
    tag_labels, tag_values = zip(*tag_counts.most_common(10))
    fig1, ax1 = plt.subplots()
    ax1.barh(tag_labels, tag_values)
    ax1.invert_yaxis()
    st.pyplot(fig1)

# Filter logic
filtered = st.session_state.transcripts
if selected_site != "All":
    filtered = [t for t in filtered if t["location"] == selected_site]
if selected_type != "All":
    filtered = [t for t in filtered if t["type"] == selected_type]
if selected_lang != "All":
    filtered = [t for t in filtered if t["language"] == selected_lang]

st.subheader("📄 Matched Results")
match_count = 0

for idx, t in enumerate(filtered):
    if not search_term or search_term.lower() in t["text"].lower():
        context_snippets = extract_context(t["text"], search_term) if search_term else [t["text"]]
        if context_snippets:
            match_count += len(context_snippets)
            with st.expander(f"{t['date']} • {t['type']} • {t['location']} • {t['language'] or 'Unknown Language'} ({len(context_snippets)} match{'es' if len(context_snippets) > 1 else ''})"):
                for snip in context_snippets:
                    highlighted = re.sub(f"({re.escape(search_term)})", r"**\\1**", snip, flags=re.IGNORECASE)
                    st.markdown(f"🔍 {highlighted}")
                    st.markdown("---")

                comment = st.text_input(f"Add comment (transcript {idx+1})", key=f"comment_{idx}")
                tag = st.text_input(f"Add tag (transcript {idx+1})", key=f"tag_{idx}")

                if st.button(f"Save comment (transcript {idx+1})"):
                    t['comments'].append(comment)
                    st.success("Comment added.")

                if st.button(f"Save tag (transcript {idx+1})"):
                    t['tags'].append(tag)
                    st.success("Tag added.")

                if t['tags']:
                    st.markdown("**Tags:** " + ", ".join(t['tags']))
                if t['comments']:
                    st.markdown("**Comments:**")
                    for c in t['comments']:
                        st.markdown(f"- {c}")

if search_term:
    st.info(f"🔎 Total matches found: {match_count}")

# Export options
st.header("📤 Export Transcripts")
col_export1, col_export2 = st.columns(2)

with col_export1:
    if st.button("Export to JSON"):
        json_data = json.dumps(st.session_state.transcripts, indent=2)
        st.download_button("Download JSON", data=json_data, file_name="transcripts.json", mime="application/json")

with col_export2:
    if st.button("Export to CSV"):
        csv_data = "text,location,type,language,date,tags,comments,themes,summary\n"
        for t in st.session_state.transcripts:
            text_escaped = t["text"].replace('"', '""')
            tags = ",".join(t["tags"])
            comments = ",".join(t["comments"])
            themes = ",".join(t.get("themes", []))
            summary = t.get("summary", "").replace('"', '""')
            csv_data += f'"{text_escaped}",{t["location"]},{t["type"]},{t["language"]},{t["date"]},"{tags}","{comments}","{themes}","{summary}"\n'
        st.download_button("Download CSV", data=csv_data, file_name="transcripts.csv", mime="text/csv")
